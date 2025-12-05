from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Настройка на стиловете
style = doc.styles['Title']
style.font.size = Pt(28)
style.font.bold = True
style.font.color.rgb = RGBColor(0, 51, 102)

style = doc.styles['Heading 1']
style.font.size = Pt(18)
style.font.bold = True
style.font.color.rgb = RGBColor(0, 102, 153)

style = doc.styles['Heading 2']
style.font.size = Pt(14)
style.font.bold = True
style.font.color.rgb = RGBColor(51, 102, 153)

# Заглавие
title = doc.add_heading('hashCode() и equals() в Java', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Подзаглавие
subtitle = doc.add_paragraph('Пълно ръководство за хаш-базираните колекции')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ============ ВЪВЕДЕНИЕ ============
doc.add_heading('1. Въведение', level=1)

intro = doc.add_paragraph()
intro.add_run('В Java, методите ').font.size = Pt(11)
intro.add_run('hashCode()').bold = True
intro.add_run(' и ').font.size = Pt(11)
intro.add_run('equals()').bold = True
intro.add_run(' са фундаментални за правилната работа на хаш-базираните колекции. '
              'Те са дефинирани в класа Object и могат да бъдат предефинирани (override) '
              'във всеки клас. Разбирането на тяхната връзка е критично важно за всеки Java разработчик.')

# ============ МЕТОД EQUALS ============
doc.add_heading('2. Методът equals()', level=1)

doc.add_heading('2.1 Какво представлява equals()?', level=2)

p = doc.add_paragraph()
p.add_run('Методът equals() се използва за сравняване на два обекта за ')
p.add_run('логическо равенство').bold = True
p.add_run('. По подразбиране, имплементацията в класа Object сравнява референциите на обектите '
          '(т.е. дали двата обекта сочат към едно и също място в паметта).')

doc.add_heading('2.2 Контракт на equals()', level=2)

p = doc.add_paragraph('Методът equals() трябва да спазва следните правила:')

# Списък с правила
rules = [
    ('Рефлексивност', 'x.equals(x) винаги връща true'),
    ('Симетричност', 'Ако x.equals(y) е true, то y.equals(x) също трябва да е true'),
    ('Транзитивност', 'Ако x.equals(y) и y.equals(z) са true, то x.equals(z) също е true'),
    ('Консистентност', 'Многократни извиквания връщат един и същ резултат (ако обектите не са променени)'),
    ('Non-null', 'x.equals(null) винаги връща false')
]

for rule, desc in rules:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{rule}: ').bold = True
    p.add_run(desc)

doc.add_heading('2.3 Пример за имплементация на equals()', level=2)

code_example = '''@Override
public boolean equals(Object obj) {
    // Проверка за същата референция
    if (this == obj) return true;
    
    // Проверка за null и тип
    if (obj == null || getClass() != obj.getClass()) 
        return false;
    
    // Кастване и сравняване на полетата
    Person person = (Person) obj;
    return age == person.age && 
           Objects.equals(name, person.name);
}'''

code_para = doc.add_paragraph()
code_para.add_run(code_example).font.name = 'Courier New'
code_para.runs[0].font.size = Pt(10)

# ============ МЕТОД HASHCODE ============
doc.add_heading('3. Методът hashCode()', level=1)

doc.add_heading('3.1 Какво представлява hashCode()?', level=2)

p = doc.add_paragraph()
p.add_run('Методът hashCode() връща ')
p.add_run('цяло число (int)').bold = True
p.add_run(', което представлява хаш кода на обекта. Този код се използва от хаш-базираните '
          'структури от данни за бързо определяне на „кофата" (bucket), в която да се съхрани или търси обектът.')

doc.add_heading('3.2 Контракт на hashCode()', level=2)

hash_rules = [
    'Ако два обекта са равни според equals(), те ТРЯБВА да имат еднакъв hashCode()',
    'Ако два обекта имат еднакъв hashCode(), те НЕ Е задължително да са равни (колизия)',
    'hashCode() трябва да връща консистентна стойност при многократни извиквания'
]

for rule in hash_rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_heading('3.3 Пример за имплементация на hashCode()', level=2)

hash_code = '''@Override
public int hashCode() {
    return Objects.hash(name, age);
}

// Или ръчна имплементация:
@Override
public int hashCode() {
    int result = 17;
    result = 31 * result + (name != null ? name.hashCode() : 0);
    result = 31 * result + age;
    return result;
}'''

code_para = doc.add_paragraph()
code_para.add_run(hash_code).font.name = 'Courier New'
code_para.runs[0].font.size = Pt(10)

# ============ ВРЪЗКА МЕЖДУ ДВАТА МЕТОДА ============
doc.add_heading('4. Връзката между equals() и hashCode()', level=1)

p = doc.add_paragraph()
p.add_run('ЗЛАТНО ПРАВИЛО: ').bold = True
p.add_run('Ако предефинирате equals(), ВИНАГИ предефинирайте и hashCode()!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Защо е толкова важно?').italic = True

reasons = [
    'Хаш-базираните колекции първо използват hashCode() за намиране на правилната „кофа"',
    'След това използват equals() за проверка на точното съвпадение',
    'Ако hashCode() не е консистентен с equals(), обектите може да „изчезнат" в колекцията'
]

for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

# ============ ХАШ-БАЗИРАНИ КОЛЕКЦИИ ============
doc.add_heading('5. Хаш-базирани колекции', level=1)

doc.add_heading('5.1 HashMap', level=2)

p = doc.add_paragraph()
p.add_run('HashMap<K, V>').bold = True
p.add_run(' е най-използваната имплементация на интерфейса Map. Тя съхранява двойки ключ-стойност.')

features = [
    'Средна времева сложност O(1) за get() и put() операции',
    'Позволява null ключове и стойности',
    'Не е thread-safe (за многонишково програмиране използвайте ConcurrentHashMap)',
    'Не гарантира подредба на елементите'
]

for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.2 HashSet', level=2)

p = doc.add_paragraph()
p.add_run('HashSet<E>').bold = True
p.add_run(' имплементира интерфейса Set и съхранява уникални елементи.')

features = [
    'Вътрешно използва HashMap (елементите са ключове, стойностите са dummy обект)',
    'Не позволява дублиращи се елементи',
    'Операциите add(), remove(), contains() са O(1) средно',
    'Позволява един null елемент'
]

for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.3 LinkedHashMap и LinkedHashSet', level=2)

p = doc.add_paragraph()
p.add_run('Тези колекции запазват ')
p.add_run('реда на добавяне').bold = True
p.add_run(' на елементите, като използват допълнителна свързана структура.')

doc.add_heading('5.4 Hashtable', level=2)

p = doc.add_paragraph()
p.add_run('Hashtable').bold = True
p.add_run(' е legacy клас, подобен на HashMap, но е synchronized (thread-safe). '
          'В модерен код се препоръчва използването на ConcurrentHashMap вместо Hashtable.')

# ============ КАК РАБОТИ ХАШИРАНЕТО ============
doc.add_heading('6. Как работи хаширането вътрешно?', level=1)

doc.add_heading('6.1 Структура на HashMap', level=2)

p = doc.add_paragraph('HashMap използва масив от „кофи" (buckets):')

steps = [
    ('Изчисляване на хаш', 'hashCode() на ключа се изчислява и трансформира до индекс в масива'),
    ('Намиране на кофа', 'Индексът определя в коя „кофа" ще се съхрани елементът'),
    ('Обработка на колизии', 'Ако две различни стойности попаднат в една кофа, се използва свързан списък или дърво'),
    ('Търсене', 'При търсене отново се изчислява хаша, намира се кофата и се търси с equals()')
]

for step, desc in steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{step}: ').bold = True
    p.add_run(desc)

doc.add_heading('6.2 Load Factor и Rehashing', level=2)

p = doc.add_paragraph()
p.add_run('Load factor').bold = True
p.add_run(' (по подразбиране 0.75) определя кога HashMap ще се преоразмери. '
          'Когато броят елементи надхвърли capacity × loadFactor, се извършва ')
p.add_run('rehashing').italic = True
p.add_run(' – създава се нов, по-голям масив и всички елементи се преразпределят.')

# ============ ЧЕСТО СРЕЩАНИ ГРЕШКИ ============
doc.add_heading('7. Често срещани грешки', level=1)

mistakes = [
    ('Предефиниране само на equals()', 
     'Обектите ще работят неправилно в HashMap/HashSet, защото ще попадат в различни кофи'),
    ('Използване на mutable полета в hashCode()', 
     'Ако полето се промени след добавяне в колекция, обектът става „невидим"'),
    ('Неконсистентна имплементация', 
     'equals() и hashCode() трябва да използват едни и същи полета'),
    ('Лош хаш алгоритъм', 
     'Може да доведе до много колизии и деградация до O(n) производителност')
]

for mistake, consequence in mistakes:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{mistake}: ').bold = True
    p.add_run(consequence)

# ============ BEST PRACTICES ============
doc.add_heading('8. Добри практики', level=1)

practices = [
    'Винаги предефинирайте hashCode() когато предефинирате equals()',
    'Използвайте Objects.hash() за лесна и надеждна имплементация',
    'Използвайте immutable полета за изчисляване на хаш кода',
    'Използвайте IDE-то за автоматично генериране на методите',
    'Пишете unit тестове за проверка на контракта',
    'Обмислете използването на record класове (Java 14+) – те автоматично генерират коректни equals() и hashCode()'
]

for practice in practices:
    doc.add_paragraph(practice, style='List Bullet')

# ============ ЗАКЛЮЧЕНИЕ ============
doc.add_heading('9. Заключение', level=1)

conclusion = doc.add_paragraph()
conclusion.add_run('Правилното разбиране и имплементация на ')
conclusion.add_run('hashCode()').bold = True
conclusion.add_run(' и ')
conclusion.add_run('equals()').bold = True
conclusion.add_run(' е от съществено значение за ефективното използване на хаш-базираните колекции в Java. '
                   'Спазването на контрактите на тези методи гарантира предвидимо поведение и оптимална '
                   'производителност на вашите приложения.')

doc.add_paragraph()

# Финален съвет
final = doc.add_paragraph()
final.add_run('💡 Запомнете: ').bold = True
final.add_run('Ако два обекта са равни (equals() връща true), '
              'техните хаш кодове ТРЯБВА да са равни. Обратното не е задължително!')

# Запазване на документа
doc.save('/workspace/HashCode_Equals_Guide.docx')
print('Документът е създаден успешно: HashCode_Equals_Guide.docx')
